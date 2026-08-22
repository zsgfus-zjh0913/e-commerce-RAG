import os
import sys
import json
import csv
import re
import hashlib

# Windows 短路径库（绕过路径长度限制安装的 torch/sentence-transformers）
_rag_libs = os.environ.get("RAG_LIBS_PATH", r"C:\rag_libs")
if os.path.isdir(_rag_libs):
    import site
    site.addsitedir(_rag_libs)
    if _rag_libs not in sys.path:
        sys.path.append(_rag_libs)
from pathlib import Path
from docx import Document

import jieba
import numpy as np
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity

os.environ.setdefault("HF_ENDPOINT", "https://hf-mirror.com")

CUSTOM_WORDS = [
    "尺码", "型号", "面料", "材质", "商品编号", "库存",
    "包邮", "退换货", "售后服务", "发货时间", "产地",
    "男装", "女装", "童装", "鞋类", "配饰", "电子产品",
    "家居用品", "食品", "美妆", "运动户外",
    "蓝牙耳机", "无线耳机", "声科达", "SC-505",
    "T恤", "卫衣", "冲锋衣", "连衣裙", "牛仔裤",
    "胸围", "腰围", "臀围", "肩宽", "衣长", "袖长",
]

for w in CUSTOM_WORDS:
    jieba.add_word(w)

SUPPORTED_EXTENSIONS = {".txt", ".md", ".csv", ".json", ".docx"}
EMBEDDING_MODEL = "BAAI/bge-small-zh-v1.5"


def tokenize(text):
    return " ".join(w.strip() for w in jieba.cut(text) if w.strip())


class RAGEngine:
    """混合检索引擎：向量语义检索 + TF-IDF 关键词检索 + 持久化索引。"""

    def __init__(self, data_dir="data", index_dir="index"):
        self.data_dir = Path(data_dir)
        self.data_dir.mkdir(parents=True, exist_ok=True)
        self.index_dir = Path(index_dir)
        self.index_dir.mkdir(parents=True, exist_ok=True)

        self._encoder = None
        self._mode = "unknown"
        self.chunks = []
        self.embeddings = None
        self.vectorizer = None
        self.tfidf_matrix = None

        self._load_or_build_index()

    # ── 嵌入模型 ─────────────────────────────────────────

    @property
    def encoder(self):
        if self._encoder is None:
            try:
                from sentence_transformers import SentenceTransformer
                print(f"[RAG] 加载嵌入模型 {EMBEDDING_MODEL} ...")
                self._encoder = SentenceTransformer(EMBEDDING_MODEL)
                self._mode = "vector"
                print("[RAG] 向量检索模式已启用")
            except Exception as e:
                print(f"[RAG] 嵌入模型加载失败: {e}")
                print("[RAG] 回退到 TF-IDF 模式")
                self._mode = "tfidf"
        return self._encoder

    # ── 持久化索引 ───────────────────────────────────────

    def _load_or_build_index(self):
        emb_path = self.index_dir / "embeddings.npy"
        chunks_path = self.index_dir / "chunks.json"
        hashes_path = self.index_dir / "doc_hashes.json"

        if chunks_path.exists():
            self.chunks = json.loads(chunks_path.read_text(encoding="utf-8"))
            if emb_path.exists() and self.embeddings is None:
                self.embeddings = np.load(str(emb_path))
                self._mode = "vector"
            print(f"[RAG] 从磁盘加载索引: {len(self.chunks)} chunks")

            if self._data_changed(hashes_path):
                print("[RAG] 检测到数据变更，重建索引...")
                self._rebuild()
            elif self.embeddings is None and self.encoder is not None:
                print("[RAG] 索引缺少向量嵌入，重建中...")
                self._rebuild()
            else:
                if self._mode != "vector":
                    self._mode = "tfidf"
                self._build_tfidf()
        else:
            print("[RAG] 首次运行，构建索引...")
            self._rebuild()

    def _data_changed(self, hashes_path):
        if not hashes_path.exists():
            return True
        old = json.loads(hashes_path.read_text(encoding="utf-8"))
        current = self._compute_doc_hashes()
        return old != current

    def _compute_doc_hashes(self):
        hashes = {}
        for f in sorted(self.data_dir.iterdir()):
            if f.is_file() and f.suffix.lower() in SUPPORTED_EXTENSIONS:
                hashes[f.name] = hashlib.md5(f.read_bytes()).hexdigest()
        return hashes

    def _save_index(self):
        if self.embeddings is not None:
            np.save(str(self.index_dir / "embeddings.npy"), self.embeddings)
        (self.index_dir / "chunks.json").write_text(
            json.dumps(self.chunks, ensure_ascii=False), encoding="utf-8"
        )
        (self.index_dir / "doc_hashes.json").write_text(
            json.dumps(self._compute_doc_hashes(), ensure_ascii=False, indent=2),
            encoding="utf-8",
        )

    # ── 文档加载 ──────────────────────────────────────────

    def _load_file(self, filepath):
        ext = filepath.suffix.lower()
        results = []

        if ext in (".txt", ".md"):
            text = filepath.read_text(encoding="utf-8")
            results.append((text, {"source": filepath.name, "type": "text"}))

        elif ext == ".csv":
            with open(filepath, "r", encoding="utf-8", newline="") as f:
                reader = csv.DictReader(f)
                for i, row in enumerate(reader):
                    text = "；".join(f"{k}：{v}" for k, v in row.items() if v)
                    results.append((text, {
                        "source": filepath.name, "type": "csv", "row": i + 2
                    }))

        elif ext == ".json":
            data = json.loads(filepath.read_text(encoding="utf-8"))
            if isinstance(data, list):
                for i, item in enumerate(data):
                    text = self._json_to_text(item)
                    results.append((text, {
                        "source": filepath.name, "type": "json", "index": i
                    }))
            elif isinstance(data, dict):
                text = self._json_to_text(data)
                results.append((text, {
                    "source": filepath.name, "type": "json"
                }))

        elif ext == ".docx":
            doc = Document(str(filepath))
            parts = []
            for para in doc.paragraphs:
                t = para.text.strip()
                if t:
                    parts.append(t)
            for table in doc.tables:
                for row in table.rows:
                    cells = [c.text.strip() for c in row.cells if c.text.strip()]
                    if cells:
                        parts.append("；".join(cells))
            text = "\n\n".join(parts)
            results.append((text, {"source": filepath.name, "type": "docx"}))

        return results

    @staticmethod
    def _json_to_text(obj, prefix=""):
        lines = []
        if isinstance(obj, dict):
            for k, v in obj.items():
                if isinstance(v, (dict, list)):
                    lines.append(f"{k}：{RAGEngine._json_to_text(v)}")
                else:
                    lines.append(f"{k}：{v}")
        elif isinstance(obj, list):
            for i, item in enumerate(obj):
                if isinstance(item, (dict, list)):
                    lines.append(f"[{i}] {RAGEngine._json_to_text(item)}")
                else:
                    lines.append(str(item))
        else:
            lines.append(str(obj))
        return "；".join(lines)

    def _split_chunks(self, text, source, meta, max_len=300):
        paragraphs = re.split(r'\n\s*\n', text.strip())
        chunks = []
        for para in paragraphs:
            para = para.strip()
            if not para:
                continue
            while len(para) > max_len:
                cut = para[:max_len]
                for sep in ['。', '；', '!', '?', '\n']:
                    pos = cut.rfind(sep)
                    if pos > max_len // 2:
                        cut = para[:pos + 1]
                        break
                chunk_meta = dict(meta)
                chunk_meta["source"] = source
                chunks.append({"text": cut, "source": source, "meta": chunk_meta})
                para = para[len(cut):]
            if para:
                chunk_meta = dict(meta)
                chunk_meta["source"] = source
                chunks.append({"text": para, "source": source, "meta": chunk_meta})
        return chunks

    # ── 索引构建 ──────────────────────────────────────────

    def _rebuild(self):
        self.chunks = []
        for filepath in sorted(self.data_dir.iterdir()):
            if filepath.is_file() and filepath.suffix.lower() in SUPPORTED_EXTENSIONS:
                try:
                    docs = self._load_file(filepath)
                    for text, meta in docs:
                        self.chunks.extend(
                            self._split_chunks(text, filepath.name, meta)
                        )
                except Exception as e:
                    print(f"[RAG] 加载 {filepath.name} 失败: {e}")

        if not self.chunks:
            print("[RAG] 数据目录无可用文档")
            self.embeddings = None
            self.vectorizer = None
            self.tfidf_matrix = None
            self._save_index()
            return

        # 向量嵌入
        if self.encoder is not None:
            texts = [c["text"] for c in self.chunks]
            self.embeddings = self.encoder.encode(
                texts, show_progress_bar=True, normalize_embeddings=True
            )
        else:
            self.embeddings = None

        # TF-IDF 辅助索引
        self._build_tfidf()
        self._save_index()
        print(f"[RAG] 索引构建完成: {len(self.chunks)} chunks, 模式={self._mode}")

    def _build_tfidf(self):
        if not self.chunks:
            self.vectorizer = None
            self.tfidf_matrix = None
            return
        corpus = [tokenize(c["text"]) for c in self.chunks]
        self.vectorizer = TfidfVectorizer(ngram_range=(1, 2))
        self.tfidf_matrix = self.vectorizer.fit_transform(corpus)

    def _add_chunks(self, new_chunks):
        """增量添加 chunks 到索引。"""
        if not new_chunks:
            return

        old_len = len(self.chunks)
        self.chunks.extend(new_chunks)

        # 增量嵌入
        if self.encoder is not None:
            texts = [c["text"] for c in new_chunks]
            new_emb = self.encoder.encode(
                texts, show_progress_bar=False, normalize_embeddings=True
            )
            if self.embeddings is not None and len(self.embeddings) > 0:
                self.embeddings = np.vstack([self.embeddings, new_emb])
            else:
                self.embeddings = new_emb

        # 重建 TF-IDF（轻量操作）
        self._build_tfidf()
        self._save_index()
        print(f"[RAG] 增量添加 {len(new_chunks)} chunks (总计 {len(self.chunks)})")

    def _remove_chunks_by_source(self, source):
        """删除指定来源的所有 chunks。"""
        old_len = len(self.chunks)
        keep_mask = [c["source"] != source for c in self.chunks]
        self.chunks = [c for i, c in enumerate(self.chunks) if keep_mask[i]]

        if self.embeddings is not None and len(self.embeddings) > 0:
            self.embeddings = self.embeddings[keep_mask]

        removed = old_len - len(self.chunks)
        if removed > 0:
            self._build_tfidf()
            self._save_index()
            print(f"[RAG] 删除 {removed} chunks (剩余 {len(self.chunks)})")
        return removed > 0

    # ── 查询扩展 ─────────────────────────────────────────

    @staticmethod
    def _expand_query(query):
        expansions = []
        height_map = {"一": "1", "二": "2", "三": "3", "四": "4", "五": "5",
                      "六": "6", "七": "7", "八": "8", "九": "9"}
        for m in re.finditer(r"一米(.)", query):
            digit = height_map.get(m.group(1))
            if digit:
                expansions.append(f"1{digit}0")

        term_map = {
            "男士": ["男装", "男"],
            "女士": ["女装", "女"],
            "多大码": ["尺码", "尺寸", "适合"],
            "穿什么": ["适合", "尺码"],
            "穿多大": ["适合", "尺码"],
            "多少钱": ["价格", "售价"],
            "包邮吗": ["运费", "包邮"],
            "能退吗": ["退换货", "退货"],
            "多久到": ["时效", "发货时间"],
            "有什么商品": ["商品", "产品"],
            "有哪些商品": ["商品", "产品"],
            "蓝牙耳机": ["耳机", "SC-505", "声科达"],
            "T恤": ["T 恤", "短袖"],
            "裤子": ["男裤", "女裤", "休闲裤", "牛仔裤"],
            "裙子": ["连衣裙"],
            "上衣": ["T 恤", "卫衣", "衬衫"],
            "冲锋衣": ["童装", "防风"],
        }
        for colloquial, formal_list in term_map.items():
            if colloquial in query:
                expansions.extend(formal_list)

        if expansions:
            return query + " " + " ".join(set(expansions))
        return query

    # ── 检索 ──────────────────────────────────────────────

    def query(self, question, top_k=5, threshold=0.0):
        if not self.chunks:
            return []

        expanded = self._expand_query(question)

        # 向量检索
        vec_scores = None
        if self.encoder is not None and self.embeddings is not None:
            q_emb = self.encoder.encode(
                [expanded], normalize_embeddings=True
            )
            vec_scores = (self.embeddings @ q_emb.T).flatten()

        # TF-IDF 检索
        tfidf_scores = None
        if self.vectorizer is not None:
            q_vec = self.vectorizer.transform([tokenize(expanded)])
            tfidf_scores = cosine_similarity(q_vec, self.tfidf_matrix).flatten()

        # 混合评分
        if vec_scores is not None and tfidf_scores is not None:
            scores = 0.7 * vec_scores + 0.3 * tfidf_scores
        elif vec_scores is not None:
            scores = vec_scores
        elif tfidf_scores is not None:
            scores = tfidf_scores
        else:
            return []

        ranked = np.argsort(scores)[::-1]
        results = []
        for idx in ranked[:top_k]:
            score = float(scores[idx])
            if score < threshold:
                continue
            chunk = self.chunks[idx]
            results.append({
                "text": chunk["text"],
                "source": chunk["source"],
                "score": round(score, 4),
                "meta": chunk.get("meta", {}),
            })
        return results

    # ── 文档管理 ──────────────────────────────────────────

    def list_documents(self):
        docs = []
        for filepath in sorted(self.data_dir.iterdir()):
            if filepath.is_file() and filepath.suffix.lower() in SUPPORTED_EXTENSIONS:
                stat = filepath.stat()
                docs.append({
                    "name": filepath.name,
                    "size": stat.st_size,
                    "modified": stat.st_mtime,
                })
        return docs

    def add_document(self, filename, content_bytes):
        filepath = self.data_dir / filename
        filepath.write_bytes(content_bytes)

        # 增量索引
        try:
            docs = self._load_file(filepath)
            new_chunks = []
            for text, meta in docs:
                new_chunks.extend(
                    self._split_chunks(text, filepath.name, meta)
                )
            self._add_chunks(new_chunks)
        except Exception as e:
            print(f"[RAG] 增量索引失败，全量重建: {e}")
            self._rebuild()
        return filepath.name

    def delete_document(self, filename):
        filepath = self.data_dir / filename
        if filepath.exists() and filepath.is_file():
            filepath.unlink()
            return self._remove_chunks_by_source(filename)
        return False

    def get_stats(self):
        return {
            "total_chunks": len(self.chunks),
            "total_documents": len(self.list_documents()),
            "mode": self._mode if self._mode != "unknown" else "tfidf",
            "index_persisted": (self.index_dir / "chunks.json").exists(),
        }
